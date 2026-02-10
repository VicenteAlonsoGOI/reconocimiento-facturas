import docx
import os

def md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"File not found: {md_path}")
        return
        
    doc = docx.Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], 0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. '):
            doc.add_paragraph(line[3:], style='List Number')
        else:
            doc.add_paragraph(line)
            
    doc.save(docx_path)
    print(f"Created: {docx_path}")

if __name__ == "__main__":
    md_to_docx("Manual_Usuario.md", "Manual_Usuario.docx")
    md_to_docx("Documentacion_Tecnica.md", "Documentacion_Tecnica.docx")
